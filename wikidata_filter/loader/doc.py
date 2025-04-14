"""
基于python-docx读取docx文件 输出文本段落及表格

{"meta": {}, "type": "paragraph", "content": "paragraph content"}
{"meta": {}, "type": "table", "content": [[]] }

@dependencies pdfminer
"""
import os
from datetime import datetime
from typing import Iterable, Any
from wikidata_filter.loader.file import BinaryFile

try:
    from docx import Document
except:
    print("Error! you need to install python-docx first: pip install python-docx")
    raise ImportError("python-docx not installed")


class Docx(BinaryFile):
    """基于python-docx读取docx文件"""
    def __init__(self, input_file: str, max_pages: int = 0, model: str = 'json', **kwargs):
        super().__init__(input_file, **kwargs)
        """
        model: 读取word文件后处理模式选择，model参数设置为json时，把word文档处理转换成统一格式的json结构数据，
        返回数据格式为{"title": "", "level": 0, "content": "", "children": []}
        model参数默认为json,传入其他参数则打印word文档段落的信息
        """
        self.max_pages = max_pages
        self.model = model

    def iter(self) -> Iterable[Any]:
        document = Document(self.instream)
        if self.model == 'json':
            json_structure = self.build_json_structure(document)
            yield json_structure
        else:
            core = document.core_properties
            props = {}
            def set_attr(k, v=None):
                if v:
                    if isinstance(v, datetime):
                        v = v.strftime('%Y-%m-%d %H:%M:%S')
                    props[k] = v
            set_attr('title', core.title)
            set_attr('subject', core.subject)
            set_attr('author', core.author)
            set_attr('keywords', core.keywords)
            set_attr('comments', core.comments)
            set_attr('created', core.created)
            set_attr('modified', core.modified)
            set_attr('language', core.language)
            set_attr('identifier', core.identifier)
            set_attr('version', core.version)
            set_attr('category', core.category)
            for paragraph in document.paragraphs:
                yield {
                    "meta": props,
                    "type": "paragraph",
                    "style": paragraph.style.name,
                    "content": paragraph.text
                }
            for table in document.tables:
                rows = [
                    [cell.text.strip() for cell in row.cells]
                    for row in table.rows
                ]
                yield {
                    "meta": props,
                    "type": "table",
                    "content": rows
                }

    def build_json_structure(self, doc):
        """
        构建JSON结构，按文档顺序处理段落和表格
        ----------
        doc: Document objec
        Returns: 结构化json数据 {"title": "", "level": 0, "content": "", "children": []}
        -------
        """
        from docx.oxml import CT_Tbl
        from docx.table import Table
        from docx.text.paragraph import Paragraph
        root = {"title": "", "level": 0, "content": "", "children": [], "tables": []}
        stack = [root]

        # 遍历文档的所有元素（段落和表格）
        for element in doc.element.body:
            if isinstance(element, CT_Tbl):
                # 处理表格
                table = Table(element, doc)
                table_data = self.parse_table(table)
                current_node = stack[-1]
                current_node["tables"].extend(table_data)
            else:
                # 处理段落
                para = Paragraph(element, doc)
                raw_text = para.text
                text = raw_text.strip() if raw_text else ""
                if text:
                    heading_text, level = self.parse_heading(para)
                    if level > 0:
                        # 处理标题节点
                        node = {"title": heading_text, "level": level, "content": "", "children": [], "tables": []}
                        while len(stack) > 1 and stack[-1]["level"] >= level:
                            stack.pop()
                        stack[-1]["children"].append(node)
                        stack.append(node)
                    else:
                        # 处理内容段落
                        if stack[-1]["content"]:
                            stack[-1]["content"] += "\n" + text
                        else:
                            stack[-1]["content"] = text

        return root["children"] if root["children"] else root

    def parse_heading(self, paragraph):
        """
        解析标题段落，返回标题文本和层级
        Parameters
        ----------
        paragraph: word段落对象

        Returns: 标题内容和几级标题（等级）
        -------
        """
        text = paragraph.text.strip() if paragraph.text else ""
        level = int(paragraph.style.name.split()[-1]) if paragraph.style.name.startswith('Heading') else 0
        return text, level

    def parse_table(self, table):
        """
        解析表格内容，返回字典列表
        Parameters
        ----------
        table: docx.table.Table object
        Returns: table list
        -------

        """
        table_data = []
        rows = table.rows
        # 提取表头（假设第一行为表头）
        if len(rows) == 0:
            return []
        header_row = rows[0]
        # 构建表格列key的列表
        headers = []
        for cell in header_row.cells:
            if cell.text:
                headers.append(cell.text.strip())
            else:
                headers.append("")
        # 处理数据行
        for row in rows[1:]:
            cells = row.cells
            row_dict = {}
            for idx, cell in enumerate(cells):
                # 如果列索引超出表头范围，跳过
                if idx < len(headers):
                    key = headers[idx]
                    value = cell.text.strip() if cell.text else ""
                    row_dict[key] = value
            table_data.append(row_dict)
        return table_data


class Doc(Docx):
    """基于libreoffice6将doc转换为docx 进而基于docx进行解析"""
    def __init__(self, input_file: str, max_pages: int = 0, **kwargs):
        super().__init__(input_file, max_pages=max_pages, auto_open=False, **kwargs)
        out_path = os.path.dirname(self.filename)
        try:
            self.doc2docx(self.filename, out_path)
        except:
            print('doc2docx failed! make sure you have install libreoffice6.4')
        self.filename = self.filename + "x"
        self.instream = open(self.filename, 'rb')

    def doc2docx(self, doc: str, output_dir: str):
        os.system(f'libreoffice6.4 --headless --convert-to docx "{doc}" --outdir "{output_dir}"')
